from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, HttpResponseBadRequest
from Crypto.Cipher import AES
from django.conf import settings
from .models import Stuff
from .forms import StuffForm
from datetime import date, datetime
import openpyxl
from docx import Document
from reportlab.pdfgen import canvas
import os, json, base64
from PIL import Image
from io import BytesIO

def create_stuff(request):
    if request.method == 'POST':
        form = StuffForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('stuff_list')
    else:
        form = StuffForm()
    return render(request, 'create_stuff.html', {'form': form})

def stuff_list(request):
    stuff = Stuff.objects.all()
    return render(request, 'stuff_list.html', {'stuff': stuff})



def home(request):
    return render(request, 'home.html')

def get_key(request):
    if request.method == "POST":
        key = request.POST.get('key')
        if len(key) != 16:
            return HttpResponseBadRequest("The key must be exactly 16 characters long.")
        return key.encode()  # Return the key as bytes
    return None

def handle_uploaded_file(uploaded_file):
    file_path = os.path.join(settings.MEDIA_ROOT, uploaded_file.name)
    with open(file_path, 'wb+') as destination:
        for chunk in uploaded_file.chunks():
            destination.write(chunk)
    return file_path

def encrypt(file_path, key):
    with open(file_path, 'rb') as file:
        plaintext = file.read()
    cipher = AES.new(key, AES.MODE_EAX)
    nonce = cipher.nonce
    ciphertext, tag = cipher.encrypt_and_digest(plaintext)
    encrypted_file_path = file_path + '.enc'
    with open(encrypted_file_path, 'wb') as file:
        file.write(nonce)
        file.write(tag)
        file.write(ciphertext)
    return encrypted_file_path

def json_serial(obj):
    """JSON serializer for objects not serializable by default json code"""
    if isinstance(obj, (datetime, date)):
        return obj.isoformat()
    raise TypeError("Type %s not serializable" % type(obj))

def encryptData(request):
    if request.method == 'POST':
        key = get_key(request)
        if isinstance(key, HttpResponseBadRequest):
            return key
        
        # Lấy tất cả dữ liệu từ mô hình Stuff
        data = list(Stuff.objects.all().values())
        plaintext = json.dumps(data, default=json_serial).encode('utf-8')
        
        # Mã hóa dữ liệu
        cipher = AES.new(key, AES.MODE_EAX)
        nonce = cipher.nonce
        ciphertext, tag = cipher.encrypt_and_digest(plaintext)
        
        # Tạo tên file và đường dẫn tự động dựa trên ngày giờ hiện tại
        new_filename = f"encrypted_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.enc"
        new_file_path = os.path.join(settings.MEDIA_ROOT, new_filename)
        
        # Lưu file mã hóa
        with open(new_file_path, 'wb') as file:
            file.write(nonce)
            file.write(tag)
            file.write(ciphertext)
        
        # Tạo response để tải file
        with open(new_file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/octet-stream')
            response['Content-Disposition'] = f'attachment; filename={new_filename}'
            return response
    return HttpResponseBadRequest("Invalid request method.")

# def display_decrypted_data(request):
#     # Trích xuất dữ liệu đã mã hóa
#     # Trả về một trang thông báo nếu không có dữ liệu nào đã được mã hóa
#     with open(encryptData(request), 'rb')

def decryptData(file_path, key, new_filename):
    with open(file_path, 'rb') as file:
        nonce = file.read(16)
        tag = file.read(16)
        ciphertext = file.read()
    cipher = AES.new(key, AES.MODE_EAX, nonce=nonce)
    plaintext = cipher.decrypt_and_verify(ciphertext, tag)
    
    data = json.loads(plaintext.decode('utf-8'))

    # Tạo tệp DOCX và chèn dữ liệu
    new_filename += '.docx'
    new_file_path = os.path.join(settings.MEDIA_ROOT, new_filename)
    doc = Document()

    for item in data:
        if 'image' in item:  # Assuming 'image' is the key for image data in base64 format
            image_data = base64.b64decode(item['image'])
            image = Image.open(BytesIO(image_data))
            image_path = os.path.join(settings.MEDIA_ROOT, 'temp_image.png')
            image.save(image_path)

            # Chèn hình ảnh vào tệp DOCX
            doc.add_picture(image_path, width=Document.DEFAULT_WIDTH/2)
            os.remove(image_path)  # Xóa ảnh sau khi chèn vào tệp DOCX

        for key, value in item.items():
            if key != 'image':  # Don't include the image data as text
                doc.add_paragraph(f"{key}: {value}")

    doc.save(new_file_path)
    return new_file_path


def decrypt(file_path, key, new_filename):
    with open(file_path, 'rb') as file:
        nonce = file.read(16)
        tag = file.read(16)
        ciphertext = file.read()
    cipher = AES.new(key, AES.MODE_EAX, nonce=nonce)
    plaintext = cipher.decrypt_and_verify(ciphertext, tag)
    
    # Ensure the new filename is secure
    new_filename = os.path.basename(new_filename)  # Sanitize the input to prevent directory traversal
    new_file_path = os.path.join(settings.MEDIA_ROOT, new_filename)
    
    with open(new_file_path, 'wb') as file:
        file.write(plaintext)
    return new_file_path

def encrypt_file(request):
    if request.method == 'POST' and 'encrypt' in request.POST:
        uploaded_file = request.FILES['file']
        file_path = handle_uploaded_file(uploaded_file)
        key = get_key(request)
        if isinstance(key, HttpResponseBadRequest):
            return key
        encrypted_file_path = encrypt(file_path, key)
        context = {'file_path': encrypted_file_path}
        return render(request, 'down.html', context)
    return render(request, 'encrypt.html')
 
def decrypt_file(request):
    if request.method == 'POST' and 'decrypt' in request.POST:
        uploaded_file = request.FILES['file']
        file_path = handle_uploaded_file(uploaded_file)
        key = get_key(request)
        if isinstance(key, HttpResponseBadRequest):
            return key
        new_filename = request.POST.get('new_filename')
        decrypted_file_path = decrypt(file_path, key, new_filename)
        context = {'file_path': decrypted_file_path}
        return render(request, 'down1.html', context)
    return render(request, 'decrypt.html')

def decrypt_file1(request):
    if request.method == 'POST' and 'decrypt' in request.POST:
        uploaded_file = request.FILES['file']
        file_path = handle_uploaded_file(uploaded_file)
        key = get_key(request)
        if isinstance(key, HttpResponseBadRequest):
            return key
        new_filename = request.POST.get('new_filename')
        decrypted_file_path = decryptData(file_path, key, new_filename)
        context = {'file_path': decrypted_file_path}
        return render(request, 'down1.html', context)
    return render(request, 'decrypt1.html')

def delete_stuff(request, stuff_id):
    stuff = get_object_or_404(Stuff, pk=stuff_id)
    stuff.delete()
    return redirect('stuff_list')
